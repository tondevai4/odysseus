import asyncio
import sqlite3 as _sql3
import email as email_mod
import email.header
import email.utils
import smtplib
import json
import re
import html
from html.parser import HTMLParser as _HTMLParser
import logging
import uuid
from datetime import datetime
from pathlib import Path
import time as _time
import threading as _threading

from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

from fastapi import APIRouter, Query, UploadFile, File, BackgroundTasks, HTTPException, Depends, Request
from fastapi.responses import FileResponse
from src.constants import DATA_DIR

from src.llm_core import llm_call_async
from src.upload_limits import read_upload_limited, EMAIL_COMPOSE_UPLOAD_MAX_BYTES

from routes.email_helpers import (
    _strip_think, _extract_reply, _apply_email_style_mechanics, require_owner, require_user, _assert_owns_account,
    _q, _attach_compose_uploads, _cleanup_compose_uploads,
    _load_settings, _save_settings, _get_email_config,
    _send_smtp_message, _smtp_security_mode,
    _IMAP_TIMEOUT_SECONDS, _open_imap_connection,
    _imap_connect, _imap, _decode_header, _detect_sent_folder, _detect_drafts_folder,
    _extract_attachment_text, _list_attachments_from_msg,
    _extract_attachment_to_disk, _extract_html, _extract_text,
    _fetch_sender_thread_context, _pre_retrieve_context,
    _EMAIL_REPLY_SYS_PROMPT_BASE, _POOL_HOOKS,
    _friendly_email_auth_error,
    SendEmailRequest, ExtractStyleRequest,
    ATTACHMENTS_DIR, COMPOSE_UPLOADS_DIR, SCHEDULED_DB,
    attachment_extract_dir, _email_cache_owner_clause,
)
from routes.email_pollers import _start_poller

logger = logging.getLogger(__name__)
ODYSSEUS_MAIL_ORIGIN = "odysseus-ui"
from .utils import *
from .state import *

router = APIRouter()
def _read_email_sync(uid, folder, account_id, owner, mark_seen=True):
    """Sync IMAP read — wrapped in to_thread by the async handler.

    Two-phase: read body in readonly to avoid races with concurrent reads
    of the same UID, then flip \\Seen in a separate readwrite session.
    BODY.PEEK[] keeps the fetch itself from tripping \\Seen.
    """
    import time as _t
    _t0 = _t.monotonic()
    raw = None
    _t_select = 0.0
    _t_fetch = 0.0
    try:
        with _imap(account_id, owner=owner) as conn:
            conn.select(_q(folder), readonly=True)
            _t_select = _t.monotonic() - _t0
            status, msg_data = _imap_uid_fetch(conn, uid, "(BODY.PEEK[])")
            _t_fetch = _t.monotonic() - _t0
            if status != "OK":
                return {"error": f"Email UID {uid} not found"}
            raw = msg_data[0][1]

        msg = email_mod.message_from_bytes(raw)

        subject = _decode_header(msg.get("Subject", "(no subject)"))
        sender = _decode_header(msg.get("From", "unknown"))
        to = _decode_header(msg.get("To", ""))
        cc = _decode_header(msg.get("Cc", ""))
        date_str = msg.get("Date", "")
        message_id = msg.get("Message-ID", "")
        in_reply_to = msg.get("In-Reply-To", "")
        references = msg.get("References", "")
        body = _extract_text(msg)
        body_html = _extract_html(msg)

        sender_name, sender_addr = email.utils.parseaddr(sender)
        parsed_date = email.utils.parsedate_to_datetime(date_str) if date_str else None
        attachments = _list_attachments_from_msg(msg)

        if mark_seen:
            # Set \Seen in a separate readwrite session so concurrent reads
            # of the same UID don't fight over a shared SELECT state.
            try:
                with _imap(account_id, owner=owner) as conn2:
                    conn2.select(_q(folder))
                    conn2.uid("STORE", _uid_bytes(uid), "+FLAGS", "\\Seen")
            except Exception:
                pass
        _t_total = _t.monotonic() - _t0
        if _t_total > 2.0:
            logger.warning(
                f"Slow email read uid={uid} folder={folder} "
                f"select={_t_select*1000:.0f}ms fetch={_t_fetch*1000:.0f}ms "
                f"size={len(raw)} total={_t_total*1000:.0f}ms"
            )

        # Look up cached summary, AI reply, and LLM-detected boundaries
        # by Message-ID
        cached_summary = None
        cached_ai_reply = None
        cached_boundaries = None
        try:
            import sqlite3 as _sql3
            _c = _sql3.connect(SCHEDULED_DB)
            owner_clause, owner_params = _email_cache_owner_clause(owner)
            _row = _c.execute(
                f"SELECT summary FROM email_summaries WHERE message_id = ? AND {owner_clause}",
                (message_id.strip(), *owner_params),
            ).fetchone()
            if _row:
                cached_summary = _row[0]
            _row2 = _c.execute(
                f"SELECT reply FROM email_ai_replies WHERE message_id = ? AND {owner_clause}",
                (message_id.strip(), *owner_params),
            ).fetchone()
            if _row2:
                cached_ai_reply = _apply_email_style_mechanics(_extract_reply(_row2[0] or ""))
            _row3 = _c.execute(
                "SELECT sig_start, quote_start, turns_json FROM email_boundaries WHERE message_id = ?",
                (message_id.strip(),),
            ).fetchone()
            cached_turns = None
            cached_sender_sig = None
            # Look up a per-sender cached signature (built by the
            # `learn_sender_signatures` action). Used by the renderer
            # to fold sigs consistently from the same address.
            try:
                if sender_addr:
                    _rs = _c.execute(
                        f"SELECT signature_text FROM sender_signatures "
                        f"WHERE from_address = ? AND {owner_clause}",
                        (sender_addr.lower().strip(), *owner_params),
                    ).fetchone()
                    if _rs and _rs[0]:
                        cached_sender_sig = _rs[0]
            except Exception:
                pass
            if _row3:
                cached_boundaries = {"sig_start": _row3[0], "quote_start": _row3[1]}
                if _row3[2]:
                    try:
                        from src.email_thread_parser import THREAD_PARSER_VERSION
                        _parsed = json.loads(_row3[2])
                        # Versioned envelope: {"v": N, "turns": [...]}.
                        # Anything else (bare list from older code, wrong
                        # version) is treated as a cache miss so the
                        # on-the-fly parser re-runs and the next write
                        # warms the cache with the current shape.
                        if (
                            isinstance(_parsed, dict)
                            and _parsed.get("v") == THREAD_PARSER_VERSION
                            and isinstance(_parsed.get("turns"), list)
                        ):
                            cached_turns = _parsed["turns"]
                    except Exception:
                        cached_turns = None
            _c.close()
        except Exception:
            pass

        # If no cached turns, parse on-the-fly so the client never has
        # to do the heavy lifting. Cheap on a 50KB body, free for short
        # ones. The background task warms the cache for next reads.
        if cached_turns is None:
            try:
                from src.email_thread_parser import parse_thread
                cached_turns = parse_thread(body_html, body)
            except Exception as _pe:
                logger.debug(f"thread parse on read failed: {_pe}")
                cached_turns = None

        return {
            "uid": uid,
            "folder": folder,
            "message_id": message_id.strip(),
            "subject": subject,
            "from_name": sender_name or sender_addr,
            "from_address": sender_addr,
            "to": to,
            "cc": cc,
            "date": parsed_date.isoformat() if parsed_date else "",
            "in_reply_to": in_reply_to.strip(),
            "references": references.strip(),
            "body": body,
            "body_html": body_html,
            "attachments": attachments,
            "cached_summary": cached_summary,
            "cached_ai_reply": cached_ai_reply,
            "boundaries": cached_boundaries,
            "thread_turns": cached_turns,
            "sender_signature": cached_sender_sig,
        }
    except Exception as e:
        logger.error(f"Failed to read email {uid}: {e}")
        return {"error": "Mail operation failed"}

def _mark_email_seen_sync(uid, folder, account_id, owner):
    try:
        with _imap(account_id, owner=owner) as conn:
            conn.select(_q(folder))
            conn.uid("STORE", _uid_bytes(uid), "+FLAGS", "\\Seen")
        _invalidate_list_cache(account_id, folder)
    except Exception as e:
        logger.debug(f"mark-seen after cached read failed uid={uid}: {e}")

@router.get("/read/{uid}")
async def read_email_by_uid(
    uid: str,
    folder: str = Query("INBOX"),
    account_id: str | None = Query(None),
    mark_seen: bool = Query(True),
    owner: str = Depends(require_owner),
):
    """Read email body. Cached for 30m, sync IMAP work runs in a thread."""
    ck = _read_cache_key(account_id, folder, uid, owner=owner)
    cached = _read_cache_get(ck)
    if cached is not None:
        if mark_seen:
            try:
                _asyncio.create_task(_asyncio.to_thread(_mark_email_seen_sync, uid, folder, account_id, owner))
            except RuntimeError:
                pass
        return cached
    result = await _asyncio.to_thread(_read_email_sync, uid, folder, account_id, owner, mark_seen)
    if result and not result.get("error"):
        _read_cache_put(ck, result)
    return result

def _schedule_recent_email_warm(emails: list, folder: str, account_id: str | None, owner: str):
    if not emails or folder == "__scheduled__":
        return
    now = _time.time()
    selected = []
    for em in emails:
        uid = str((em or {}).get("uid") or "").strip()
        if not uid:
            continue
        try:
            epoch = float((em or {}).get("date_epoch") or 0)
        except Exception:
            epoch = 0
        if epoch and now - epoch > _WARM_RECENT_SECONDS:
            continue
        try:
            size = int((em or {}).get("size") or 0)
        except Exception:
            size = 0
        if size > _WARM_MAX_BYTES:
            continue
        ck = _read_cache_key(account_id, folder, uid, owner=owner)
        if _read_cache_get(ck) is not None or ck in _WARMING_READS:
            continue
        _WARMING_READS.add(ck)
        selected.append((uid, ck))
        if len(selected) >= _WARM_READ_LIMIT:
            break
    if not selected:
        return

    async def _warm():
        for uid, ck in selected:
            if _read_cache_get(ck) is not None:
                _WARMING_READS.discard(ck)
                continue
            try:
                result = await _asyncio.to_thread(_read_email_sync, uid, folder, account_id, owner, False)
                if result and not result.get("error"):
                    _read_cache_put(ck, result)
            except Exception as e:
                logger.debug(f"email read warm skipped uid={uid}: {e}")
            finally:
                _WARMING_READS.discard(ck)
                await _asyncio.sleep(0.05)

    try:
        _asyncio.create_task(_warm())
    except RuntimeError:
        pass

@router.get("/attachments/{uid}")
async def list_attachments(uid: str, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
    """List attachments for an email."""
    try:
        with _imap(account_id, owner=owner) as conn:
            conn.select(_q(folder), readonly=True)
            status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
        if status != "OK":
            return {"attachments": [], "error": "Email not found"}
        raw = msg_data[0][1]
        msg = email_mod.message_from_bytes(raw)
        attachments = _list_attachments_from_msg(msg)
        return {"attachments": attachments, "uid": uid}
    except Exception as e:
        logger.error(f"Failed to list attachments for {uid}: {e}")
        return {"attachments": [], "error": "Mail operation failed"}

@router.get("/attachment/{uid}/{index}")
async def download_attachment(uid: str, index: int, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
    """Download a specific attachment by email UID and attachment index. Saves to local disk and returns the file."""
    try:
        with _imap(account_id, owner=owner) as conn:
            conn.select(_q(folder), readonly=True)
            status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
        if status != "OK":
            return {"error": "Email not found"}
        raw = msg_data[0][1]
        msg = email_mod.message_from_bytes(raw)

        # Extract to a per-email folder
        target_dir = attachment_extract_dir(folder, uid)
        filepath = _extract_attachment_to_disk(msg, index, target_dir)
        if not filepath:
            return {"error": f"Attachment index {index} not found"}

        return FileResponse(
            path=str(filepath),
            filename=filepath.name,
            media_type="application/octet-stream",
        )
    except Exception as e:
        logger.error(f"Failed to download attachment {uid}/{index}: {e}")
        return {"error": "Mail operation failed"}

@router.post("/attachment-as-doc/{uid}/{index}")
async def attachment_as_doc(uid: str, index: int, request: Request, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
    """Extract an email attachment and open it in the document editor.

    Supported extensions:
      - .pdf   → rendered as PDF Document (existing flow)
      - .docx  → text extracted to markdown Document
      - .txt / .md → loaded directly as a markdown Document

    Returns {doc_id} so the frontend can open it as a tab in the doc panel.
    Other types are rejected — caller should fall back to download.
    """
    try:
        with _imap(account_id, owner=owner) as conn:
            conn.select(_q(folder), readonly=True)
            status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
        if status != "OK":
            return {"error": "Email not found"}
        raw = msg_data[0][1]
        msg = email_mod.message_from_bytes(raw)

        target_dir = attachment_extract_dir(folder, uid)
        filepath = _extract_attachment_to_disk(msg, index, target_dir)
        if not filepath:
            return {"error": f"Attachment index {index} not found"}

        from pathlib import Path as _Path
        base = _Path(filepath).name
        if base.startswith("."):
            return {"error": "Invalid filename", "filename": base}
        ext = _Path(base).suffix.lower()

        import os as _os
        title = _os.path.splitext(filepath.name)[0]

        # Capture the source email's identity so the doc can later be used
        # to thread a signed-reply back to the original sender.
        src_message_id = (msg.get("Message-ID") or "").strip()
        def _tag_doc_with_source(doc_id_to_tag: str):
            if not doc_id_to_tag:
                return
            try:
                from src.database import SessionLocal as _SL, Document as _Doc
                _db = _SL()
                try:
                    d = _db.query(_Doc).filter(_Doc.id == doc_id_to_tag).first()
                    if d:
                        d.source_email_uid = str(uid)
                        d.source_email_folder = folder
                        d.source_email_account_id = account_id or ""
                        d.source_email_message_id = src_message_id
                        _db.commit()
                finally:
                    _db.close()
            except Exception as _e:
                logger.warning(f"tag doc source-email failed: {_e}")

        # Extracted docs MUST belong to a session the caller owns — a
        # session-less ("orphan") doc is rejected by get_document's owner
        # check (404), so the frontend's loadDocument() throws and nothing
        # opens (the "open in document didn't open" bug). Attach it to the
        # user's most-recent session so it's fetchable + ownable.
        from src.auth_helpers import get_current_user as _gcu
        _doc_user = _gcu(request)
        def _resolve_doc_session():
            try:
                from src.database import SessionLocal as _SL, Session as _Sess
                _db = _SL()
                try:
                    _q2 = _db.query(_Sess)
                    if _doc_user:
                        _q2 = _q2.filter(_Sess.owner == _doc_user)
                    s = _q2.order_by(_Sess.updated_at.desc()).first()
                    return s.id if s else None
                finally:
                    _db.close()
            except Exception as _e:
                logger.warning(f"resolve doc session failed: {_e}")
                return None
        doc_session_id = _resolve_doc_session()

        # ── PDF path (existing) ────────────────────────────────────
        if ext == ".pdf":
            import shutil as _shutil
            from src.constants import UPLOAD_DIR
            from src.pdf_forms import has_form_fields, extract_fields
            from src.pdf_form_doc import (
                save_field_sidecar,
                create_form_markdown_document,
                create_plain_pdf_document,
            )

            upload_id = f"{uuid.uuid4().hex}.pdf"
            today = datetime.utcnow().strftime("%Y/%m/%d")
            dated_dir = _os.path.join(UPLOAD_DIR, today)
            _os.makedirs(dated_dir, exist_ok=True)
            dest_path = _os.path.join(dated_dir, upload_id)
            _shutil.copyfile(str(filepath), dest_path)

            is_form = False
            try:
                is_form = has_form_fields(dest_path)
            except Exception as e:
                logger.warning(f"has_form_fields failed for attachment PDF: {e}")

            if is_form:
                fields = extract_fields(dest_path)
                save_field_sidecar(dest_path, fields)
                doc_id = create_form_markdown_document(
                    session_id=doc_session_id,
                    fields=fields,
                    upload_id=upload_id,
                    title=title,
                    intro_text=None,
                )
            else:
                doc_id = create_plain_pdf_document(
                    session_id=doc_session_id,
                    upload_id=upload_id,
                    title=title,
                )

            if not doc_id:
                return {"error": "Failed to create document"}
            _tag_doc_with_source(doc_id)
            return {"doc_id": doc_id, "filename": filepath.name}

        # ── DOCX path: extract text → markdown document ───────────
        if ext == ".docx":
            try:
                from docx import Document as _Docx
            except ImportError:
                return {"error": "python-docx not installed", "filename": base}
            try:
                d = _Docx(str(filepath))
            except Exception as e:
                return {"error": f"Failed to read docx: {e}", "filename": base}
            # Convert paragraphs to markdown — preserve heading styles as #/##/###,
            # bullet lists as `- `, numbered lists as `1.`, and keep tables as
            # simple pipe-delimited rows.
            lines: list[str] = []
            for p in d.paragraphs:
                text = p.text or ""
                style = (p.style.name if p.style else "") or ""
                if not text.strip():
                    lines.append("")
                    continue
                if style.startswith("Heading 1"): lines.append(f"# {text}")
                elif style.startswith("Heading 2"): lines.append(f"## {text}")
                elif style.startswith("Heading 3"): lines.append(f"### {text}")
                elif style.startswith("Heading "): lines.append(f"#### {text}")
                elif style.startswith("List Bullet"): lines.append(f"- {text}")
                elif style.startswith("List Number"): lines.append(f"1. {text}")
                else: lines.append(text)
            for tbl in d.tables:
                lines.append("")
                for ri, row in enumerate(tbl.rows):
                    cells = [(c.text or "").replace("|", "\\|").replace("\n", " ").strip() for c in row.cells]
                    lines.append("| " + " | ".join(cells) + " |")
                    if ri == 0:
                        lines.append("|" + "|".join(["---"] * len(cells)) + "|")
                lines.append("")
            content = "\n".join(lines).strip() or f"_(empty {base})_"

            from src.database import SessionLocal as _SL, Document as _Doc, DocumentVersion as _DV
            doc_id = str(uuid.uuid4())
            ver_id = str(uuid.uuid4())
            _db = _SL()
            try:
                _db.query(_Doc).filter(_Doc.is_active == True).update({"is_active": False})
                _db.add(_Doc(
                    id=doc_id, session_id=doc_session_id, title=title,
                    language="markdown", current_content=content,
                    version_count=1, is_active=True,
                ))
                _db.add(_DV(
                    id=ver_id, document_id=doc_id, version_number=1,
                    content=content, summary="Imported from DOCX", source="upload",
                ))
                _db.commit()
            finally:
                _db.close()
            _tag_doc_with_source(doc_id)
            return {"doc_id": doc_id, "filename": filepath.name}

        # ── Plain text / markdown ────────────────────────────────
        if ext in (".txt", ".md", ".markdown"):
            try:
                content = filepath.read_text(encoding="utf-8", errors="replace")
            except Exception as e:
                return {"error": f"Failed to read text file: {e}", "filename": base}
            from src.database import SessionLocal as _SL, Document as _Doc, DocumentVersion as _DV
            doc_id = str(uuid.uuid4())
            ver_id = str(uuid.uuid4())
            _db = _SL()
            try:
                _db.query(_Doc).filter(_Doc.is_active == True).update({"is_active": False})
                _db.add(_Doc(
                    id=doc_id, session_id=doc_session_id, title=title,
                    language="markdown", current_content=content,
                    version_count=1, is_active=True,
                ))
                _db.add(_DV(
                    id=ver_id, document_id=doc_id, version_number=1,
                    content=content, summary="Imported from email attachment", source="upload",
                ))
                _db.commit()
            finally:
                _db.close()
            _tag_doc_with_source(doc_id)
            return {"doc_id": doc_id, "filename": filepath.name}

        return {"error": f"Unsupported attachment type: {ext}", "filename": base}
    except Exception as e:
        logger.error(f"attachment-as-doc {uid}/{index} failed: {e}")
        return {"error": "Mail operation failed"}

@router.post("/attachment-path/{uid}/{index}")
async def get_attachment_path(uid: str, index: int, folder: str = Query("INBOX"), account_id: str | None = Query(None), owner: str = Depends(require_owner)):
    """Extract attachment to local disk and return the path (for AI to read via read_file)."""
    try:
        with _imap(account_id, owner=owner) as conn:
            conn.select(_q(folder), readonly=True)
            status, msg_data = _imap_uid_fetch(conn, uid, "(RFC822)")
        if status != "OK":
            return {"error": "Email not found"}
        raw = msg_data[0][1]
        msg = email_mod.message_from_bytes(raw)

        target_dir = attachment_extract_dir(folder, uid)
        filepath = _extract_attachment_to_disk(msg, index, target_dir)
        if not filepath:
            return {"error": f"Attachment index {index} not found"}

        return {"path": str(filepath), "filename": filepath.name, "size": filepath.stat().st_size}
    except Exception as e:
        logger.error(f"Failed to get attachment path {uid}/{index}: {e}")
        return {"error": "Mail operation failed"}

